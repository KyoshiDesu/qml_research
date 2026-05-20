from __future__ import annotations

import csv
import json
import math
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "assets"
DATA_DIR = OUT_DIR / "data"
REPORT_PATH = OUT_DIR / "qml_experiment_report.docx"
MARKDOWN_PATH = OUT_DIR / "qml_experiment_report.md"
SUMMARY_PATH = DATA_DIR / "summary.json"

DATASETS = [
    ("iris", ROOT / "outputs_iris"),
    ("wine", ROOT / "outputs_wine"),
    ("heart", ROOT / "outputs_heart"),
    ("breast", ROOT / "outputs_breast"),
]

COLORS = {
    "bg": "#F6F1E8",
    "ink": "#172033",
    "muted": "#5A6476",
    "rule": "#D8CBB6",
    "svm": "#1F5EFF",
    "qsvm": "#F06C3B",
    "accent": "#0E766E",
    "warn": "#9A3D1A",
    "panel": "#FFFDF9",
}


@dataclass
class DatasetSummary:
    key: str
    display_name: str
    source: str
    rows: int
    features: int
    classes: int
    svm_accuracy: float
    qsvm_accuracy: float
    svm_macro_f1: float
    qsvm_macro_f1: float
    classical_search_seconds: float
    qsvm_search_seconds: float
    qsvm_confirm_seconds: float
    qsvm_final_train_seconds: float
    total_classical_seconds: float
    total_qsvm_seconds: float
    gpu_available: bool
    gpu_used: bool
    warnings: list[str]
    log_highlights: list[str]
    confusion_note: str
    best_classical_params: str
    best_qsvm_params: str


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    DATA_DIR.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def fmt_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def fmt_sec(value: float) -> str:
    if value >= 3600:
        return f"{value / 3600:.2f} h"
    if value >= 60:
        return f"{value / 60:.1f} min"
    return f"{value:.1f} s"


def read_log_highlights(log_text: str) -> list[str]:
    highlights: list[str] = []
    for line in log_text.splitlines():
        if "ERROR" in line:
            highlights.append(line.strip())
        elif "WARNING" in line and "QSVM Optuna trial" in line:
            highlights.append(line.strip())
        elif "Experiment completed successfully." in line:
            highlights.append(line.strip())
    return highlights[:8]


def build_confusion_note(error_analysis: dict, dataset_key: str) -> str:
    binary_terms = error_analysis.get("binary_error_terms")
    if binary_terms:
        return (
            f"Type I errors: {binary_terms['false_positive_type_i_error']}, "
            f"Type II errors: {binary_terms['false_negative_type_ii_error']} "
            f"for positive class {binary_terms['positive_label']}."
        )
    pairs = error_analysis.get("most_confused_class_pairs") or []
    if not pairs:
        return f"No dominant misclassification pair surfaced for {dataset_key}."
    pair = pairs[0]
    return f"Most frequent confusion: {pair['actual']} predicted as {pair['predicted']} ({pair['count']} case(s))."


def summarize_dataset(key: str, folder: Path) -> DatasetSummary:
    dataset = load_json(folder / "metadata" / "dataset_summary.json")
    svm_metrics = load_json(folder / "metadata" / "svm_test_metrics.json")
    qsvm_metrics = load_json(folder / "metadata" / "qsvm_test_metrics.json")
    resource = load_json(folder / "metadata" / "resource_summary.json")
    compute = load_json(folder / "metadata" / "compute_allocation_summary.json")
    svm_error = load_json(folder / "metadata" / "svm_error_analysis.json")
    qsvm_info = load_json(folder / "metadata" / "qsvm_best_confirmed_model_info.json")
    svm_params = load_json(folder / "metadata" / "svm_best_params.json")
    log_text = (folder / "logs" / "experiment.log").read_text(encoding="utf-8")

    phases = resource["phase_summaries"]
    classical_total = (
        phases["preprocessing"]["duration_seconds"]
        + phases["classical_svm_grid_search"]["duration_seconds"]
        + phases["classical_svm_final_training"]["duration_seconds"]
    )
    qsvm_total = (
        phases["qsvm_optuna_search"]["duration_seconds"]
        + phases["qsvm_confirmation_phase"]["duration_seconds"]
        + phases["qsvm_final_training"]["duration_seconds"]
    )

    best_qsvm_params = qsvm_info["best_confirmed_params"]
    best_qsvm_text = (
        f"{best_qsvm_params['feature_map_type']}, reps={best_qsvm_params['reps']}, "
        f"entanglement={best_qsvm_params['entanglement']}, paulis={','.join(best_qsvm_params['paulis'])}"
    )

    display_name = {
        "iris": "Iris",
        "wine": "Wine",
        "heart": "Heart Disease",
        "breast": "Breast Cancer",
    }[key]

    return DatasetSummary(
        key=key,
        display_name=display_name,
        source=dataset["dataset_source"],
        rows=dataset["row_count_after_target_drop"],
        features=dataset["feature_column_count"],
        classes=dataset["class_count"],
        svm_accuracy=svm_metrics["accuracy"],
        qsvm_accuracy=qsvm_metrics["accuracy"],
        svm_macro_f1=svm_metrics["macro_f1"],
        qsvm_macro_f1=qsvm_metrics["macro_f1"],
        classical_search_seconds=phases["classical_svm_grid_search"]["duration_seconds"],
        qsvm_search_seconds=phases["qsvm_optuna_search"]["duration_seconds"],
        qsvm_confirm_seconds=phases["qsvm_confirmation_phase"]["duration_seconds"],
        qsvm_final_train_seconds=phases["qsvm_final_training"]["duration_seconds"],
        total_classical_seconds=classical_total,
        total_qsvm_seconds=qsvm_total,
        gpu_available=resource["gpu_available"],
        gpu_used=compute["gpu_used"],
        warnings=resource.get("warnings", []),
        log_highlights=read_log_highlights(log_text),
        confusion_note=build_confusion_note(svm_error, key),
        best_classical_params=", ".join(f"{k}={v}" for k, v in svm_params.items()),
        best_qsvm_params=best_qsvm_text,
    )


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/aptos.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
    ]
    bold_candidates = [
        "C:/Windows/Fonts/aptos-bold.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
        "C:/Windows/Fonts/calibrib.ttf",
    ]
    paths = bold_candidates if bold else candidates
    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_bar_chart(datasets: list[DatasetSummary]) -> Path:
    width, height = 1400, 860
    image = Image.new("RGB", (width, height), COLORS["bg"])
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, True)
    body_font = get_font(20)
    small_font = get_font(18)
    bold_font = get_font(20, True)

    draw.text((70, 48), "Accuracy by Dataset: Classical SVM vs QSVM", fill=COLORS["ink"], font=title_font)
    draw.text((70, 98), "The only tie is Iris; every other dataset favors the classical baseline.", fill=COLORS["muted"], font=body_font)

    chart_left, chart_top = 90, 190
    chart_right, chart_bottom = 1320, 700
    draw.rectangle((chart_left, chart_top, chart_right, chart_bottom), outline=COLORS["rule"], width=2, fill=COLORS["panel"])

    max_value = 1.0
    for tick in range(0, 11):
        y = chart_bottom - (chart_bottom - chart_top - 40) * tick / 10
        draw.line((chart_left + 70, y, chart_right - 30, y), fill="#E7DFD1", width=1)
        draw.text((chart_left + 10, y - 10), f"{tick * 10}%", fill=COLORS["muted"], font=small_font)

    usable_width = chart_right - chart_left - 120
    group_width = usable_width / len(datasets)
    bar_width = 72
    inner_gap = 18
    usable_height = chart_bottom - chart_top - 70

    for idx, ds in enumerate(datasets):
        group_x = chart_left + 95 + idx * group_width
        svm_h = usable_height * ds.svm_accuracy / max_value
        qsvm_h = usable_height * ds.qsvm_accuracy / max_value
        x1 = group_x
        x2 = group_x + bar_width + inner_gap
        y1 = chart_bottom - svm_h
        y2 = chart_bottom - qsvm_h

        draw.rounded_rectangle((x1, y1, x1 + bar_width, chart_bottom), radius=14, fill=COLORS["svm"])
        draw.rounded_rectangle((x2, y2, x2 + bar_width, chart_bottom), radius=14, fill=COLORS["qsvm"])
        draw.text((x1 - 2, y1 - 28), fmt_pct(ds.svm_accuracy), fill=COLORS["ink"], font=bold_font)
        draw.text((x2 - 2, y2 - 28), fmt_pct(ds.qsvm_accuracy), fill=COLORS["ink"], font=bold_font)
        label_bbox = draw.textbbox((0, 0), ds.display_name, font=bold_font)
        label_x = group_x + 40 - (label_bbox[2] - label_bbox[0]) / 2
        draw.text((label_x, chart_bottom + 18), ds.display_name, fill=COLORS["ink"], font=bold_font)
        delta = ds.svm_accuracy - ds.qsvm_accuracy
        delta_text = "Tie" if abs(delta) < 0.0001 else f"SVM +{delta * 100:.1f} pts"
        delta_fill = COLORS["accent"] if delta >= 0 else COLORS["warn"]
        draw.text((group_x - 10, chart_bottom + 48), delta_text, fill=delta_fill, font=small_font)

    legend_y = 770
    draw.rounded_rectangle((94, legend_y, 122, legend_y + 28), radius=8, fill=COLORS["svm"])
    draw.text((136, legend_y + 2), "Classical SVM", fill=COLORS["ink"], font=body_font)
    draw.rounded_rectangle((330, legend_y, 358, legend_y + 28), radius=8, fill=COLORS["qsvm"])
    draw.text((372, legend_y + 2), "QSVM", fill=COLORS["ink"], font=body_font)

    out_path = ASSET_DIR / "accuracy_comparison.png"
    image.save(out_path)
    return out_path


def draw_runtime_chart(datasets: list[DatasetSummary]) -> Path:
    width, height = 1400, 900
    image = Image.new("RGB", (width, height), COLORS["bg"])
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, True)
    body_font = get_font(20)
    small_font = get_font(18)
    bold_font = get_font(20, True)

    draw.text((70, 48), "Runtime Burden: QSVM Search Dominates the Workflow", fill=COLORS["ink"], font=title_font)
    draw.text((70, 98), "Bars show end-to-end model-building time excluding visualization; QSVM remains two to four orders of magnitude slower.", fill=COLORS["muted"], font=body_font)

    chart_left, chart_top = 90, 190
    chart_right, chart_bottom = 1320, 710
    draw.rectangle((chart_left, chart_top, chart_right, chart_bottom), outline=COLORS["rule"], width=2, fill=COLORS["panel"])

    max_value = max(ds.total_qsvm_seconds for ds in datasets)
    usable_width = chart_right - chart_left - 120
    group_width = usable_width / len(datasets)
    bar_width = 80
    usable_height = chart_bottom - chart_top - 70

    for tick in range(0, 6):
        ratio = tick / 5
        y = chart_bottom - usable_height * ratio
        draw.line((chart_left + 80, y, chart_right - 30, y), fill="#E7DFD1", width=1)
        label = fmt_sec(max_value * ratio)
        draw.text((chart_left + 10, y - 10), label, fill=COLORS["muted"], font=small_font)

    for idx, ds in enumerate(datasets):
        group_x = chart_left + 120 + idx * group_width
        classical_h = usable_height * ds.total_classical_seconds / max_value
        qsvm_h = usable_height * ds.total_qsvm_seconds / max_value
        x1 = group_x
        x2 = group_x + bar_width + 24
        y1 = chart_bottom - classical_h
        y2 = chart_bottom - qsvm_h

        draw.rounded_rectangle((x1, y1, x1 + bar_width, chart_bottom), radius=14, fill=COLORS["svm"])
        draw.rounded_rectangle((x2, y2, x2 + bar_width, chart_bottom), radius=14, fill=COLORS["qsvm"])
        draw.text((x1 - 8, y1 - 26), fmt_sec(ds.total_classical_seconds), fill=COLORS["ink"], font=small_font)
        draw.text((x2 - 8, y2 - 26), fmt_sec(ds.total_qsvm_seconds), fill=COLORS["ink"], font=small_font)

        label_bbox = draw.textbbox((0, 0), ds.display_name, font=bold_font)
        label_x = group_x + 52 - (label_bbox[2] - label_bbox[0]) / 2
        draw.text((label_x, chart_bottom + 18), ds.display_name, fill=COLORS["ink"], font=bold_font)

        ratio = ds.total_qsvm_seconds / max(ds.total_classical_seconds, 0.001)
        draw.text((group_x - 18, chart_bottom + 48), f"{ratio:.0f}x slower", fill=COLORS["warn"], font=small_font)

    out_path = ASSET_DIR / "runtime_comparison.png"
    image.save(out_path)
    return out_path


def draw_log_findings_card(findings: list[str]) -> Path:
    width, height = 1400, 820
    image = Image.new("RGB", (width, height), COLORS["bg"])
    draw = ImageDraw.Draw(image)
    title_font = get_font(34, True)
    body_font = get_font(22)
    small_font = get_font(18)

    draw.text((70, 48), "Log-Derived Findings", fill=COLORS["ink"], font=title_font)
    draw.text((70, 98), "Operational patterns that showed up repeatedly across experiment logs.", fill=COLORS["muted"], font=body_font)

    y = 180
    for idx, finding in enumerate(findings, start=1):
        draw.rounded_rectangle((80, y, 1320, y + 110), radius=18, fill=COLORS["panel"], outline=COLORS["rule"], width=2)
        draw.rounded_rectangle((104, y + 28, 148, y + 72), radius=12, fill=COLORS["ink"])
        draw.text((118, y + 34), f"{idx}", fill="#FFFFFF", font=body_font)
        lines = wrap_text(draw, finding, body_font, 1120)
        yy = y + 24
        for line in lines[:3]:
            draw.text((182, yy), line, fill=COLORS["ink"], font=body_font)
            yy += 28
        y += 134

    draw.text((80, 754), "Source: experiment logs and resource summaries in each dataset folder.", fill=COLORS["muted"], font=small_font)
    out_path = ASSET_DIR / "log_findings.png"
    image.save(out_path)
    return out_path


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 32, 51)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181) if style_name != "Heading 3" else RGBColor(31, 77, 120)


def add_table_after(doc: Document, datasets: list[DatasetSummary]) -> None:
    doc.add_heading("Cross-Dataset Results Summary", level=1)
    doc.add_paragraph(
        "This comparison uses the final test metrics and phase-level runtime summaries stored inside each dataset output folder."
    )
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Dataset", "Rows", "Features", "SVM Acc.", "QSVM Acc.", "Gap", "QSVM Build Time"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "F2F4F7")

    for ds in datasets:
        row = table.add_row().cells
        delta = ds.svm_accuracy - ds.qsvm_accuracy
        values = [
            ds.display_name,
            str(ds.rows),
            str(ds.features),
            fmt_pct(ds.svm_accuracy),
            fmt_pct(ds.qsvm_accuracy),
            "Tie" if abs(delta) < 0.0001 else f"{delta * 100:.1f} pts",
            fmt_sec(ds.total_qsvm_seconds),
        ]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)

    doc.add_picture(str(ASSET_DIR / "accuracy_comparison.png"), width=Inches(6.4))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_runtime_section(doc: Document, datasets: list[DatasetSummary]) -> None:
    doc.add_heading("Compute and Runtime Findings", level=1)
    doc.add_paragraph(
        "Across every dataset, the heaviest phase was the QSVM Optuna search. Classical model selection stayed under three seconds, while QSVM build time ranged from about 12 minutes on Iris to more than 2.4 hours on Breast Cancer."
    )
    doc.add_picture(str(ASSET_DIR / "runtime_comparison.png"), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bullet_points = [
        "QSVM search plus confirmation dominated total runtime in all four experiments.",
        "CPU saturation warnings appeared in every dataset, suggesting the pipeline was compute-bound rather than I/O-bound.",
        "The Heart Disease run had GPU visibility, but the resource summary still flagged the GPU as mostly idle in several phases.",
        "No dataset reported sustained memory pressure, so runtime inefficiency appears tied to search cost rather than RAM exhaustion.",
    ]
    for bullet in bullet_points:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)


def add_dataset_findings(doc: Document, datasets: list[DatasetSummary]) -> None:
    doc.add_heading("Dataset-Specific Findings", level=1)
    for ds in datasets:
        delta = ds.svm_accuracy - ds.qsvm_accuracy
        doc.add_heading(ds.display_name, level=2)
        lead = (
            f"Classical SVM reached {fmt_pct(ds.svm_accuracy)} accuracy versus "
            f"{fmt_pct(ds.qsvm_accuracy)} for QSVM."
        )
        if abs(delta) < 0.0001:
            lead += " This was the only dataset where both models tied."
        else:
            lead += f" The classical model led by {delta * 100:.1f} percentage points."
        doc.add_paragraph(lead)
        detail_points = [
            f"Dataset shape: {ds.rows} rows, {ds.features} features, {ds.classes} classes; source={ds.source}.",
            f"Best classical settings: {ds.best_classical_params}.",
            f"Best confirmed QSVM settings: {ds.best_qsvm_params}.",
            ds.confusion_note,
        ]
        for point in detail_points:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(point)


def add_log_section(doc: Document, datasets: list[DatasetSummary]) -> None:
    doc.add_heading("Log Review", level=1)
    doc.add_paragraph(
        "The logs add useful operational context beyond the final metrics. One configuration bug surfaced early, and the QSVM search repeatedly recorded failed Optuna trials before the runs still completed successfully."
    )
    doc.add_picture(str(ASSET_DIR / "log_findings.png"), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ds in datasets:
        doc.add_heading(f"{ds.display_name} log notes", level=2)
        for line in ds.log_highlights[:4]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line)


def add_recommendations(doc: Document) -> None:
    doc.add_heading("Recommendations", level=1)
    for item in [
        "Keep the classical SVM as the primary baseline and likely production default for this workflow.",
        "Treat QSVM as an exploratory model class unless there is a clear quantum-specific objective beyond accuracy, because the current runtime cost is disproportionate to the observed gains.",
        "Fix or investigate the repeated Optuna trial failures (`'NoneType' object has no attribute 'items'`) before running larger studies.",
        "Standardize the experiment environment, because Heart Disease was executed on a materially stronger machine than the other runs, which complicates direct runtime comparisons.",
        "If QSVM work continues, prioritize reducing search space and profiling the feature-map evaluation path before scaling to larger datasets.",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def build_docx(datasets: list[DatasetSummary]) -> None:
    doc = Document()
    format_doc_styles(doc)

    doc.add_paragraph("QML Research", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("Cross-dataset report for the combined experiment outputs")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(90, 100, 118)

    doc.add_paragraph(
        "This report synthesizes the logs, metrics, plots, and metadata stored under `combined_outputs` for Iris, Wine, Heart Disease, and Breast Cancer."
    )

    doc.add_heading("Executive Summary", level=1)
    summary_points = [
        "Classical SVM matched or outperformed QSVM on all four datasets.",
        "Iris was the only tie; Wine, Heart Disease, and Breast Cancer all favored the classical model.",
        "QSVM search and confirmation phases dominated runtime by a wide margin in every run.",
        "Operational logs showed one early parameter-grid setup error and repeated QSVM Optuna trial failures, but all final runs completed successfully.",
    ]
    for point in summary_points:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(point)

    add_table_after(doc, datasets)
    add_runtime_section(doc, datasets)
    add_dataset_findings(doc, datasets)
    add_log_section(doc, datasets)
    add_recommendations(doc)

    doc.save(REPORT_PATH)


def build_markdown(datasets: list[DatasetSummary]) -> None:
    lines: list[str] = []
    lines.append("# QML Experiment Report")
    lines.append("")
    lines.append("## Executive Summary")
    lines.append("")
    lines.append("- Classical SVM matched or beat QSVM on every dataset.")
    lines.append("- Iris was a tie at 93.3% accuracy for both models.")
    lines.append("- Classical SVM led on Wine (+5.6 pts), Heart Disease (+1.6 pts), and Breast Cancer (+2.6 pts).")
    lines.append("- QSVM build time was dramatically larger in every case because Optuna search and confirmation dominated runtime.")
    lines.append("")
    lines.append("## Results Table")
    lines.append("")
    lines.append("| Dataset | Rows | Features | SVM Accuracy | QSVM Accuracy | Gap | QSVM Build Time |")
    lines.append("|---|---:|---:|---:|---:|---:|---:|")
    for ds in datasets:
        delta = ds.svm_accuracy - ds.qsvm_accuracy
        gap_text = "Tie" if abs(delta) < 0.0001 else f"{delta * 100:.1f} pts"
        lines.append(
            f"| {ds.display_name} | {ds.rows} | {ds.features} | {fmt_pct(ds.svm_accuracy)} | {fmt_pct(ds.qsvm_accuracy)} | {gap_text} | {fmt_sec(ds.total_qsvm_seconds)} |"
        )
    lines.append("")
    lines.append("## Recommendations")
    lines.append("")
    for rec in [
        "Keep classical SVM as the main benchmark and likely default model for this pipeline.",
        "Investigate the repeated QSVM Optuna trial failures before running broader studies.",
        "Normalize hardware conditions if future runtime comparisons are meant to be strict apples-to-apples comparisons.",
    ]:
        lines.append(f"- {rec}")
    lines.append("")
    MARKDOWN_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_summary_json(datasets: list[DatasetSummary]) -> None:
    ranked = sorted(datasets, key=lambda d: d.svm_accuracy - d.qsvm_accuracy, reverse=True)
    summary = {
        "datasets": [
            {
                "key": ds.key,
                "display_name": ds.display_name,
                "rows": ds.rows,
                "features": ds.features,
                "classes": ds.classes,
                "svm_accuracy": ds.svm_accuracy,
                "qsvm_accuracy": ds.qsvm_accuracy,
                "svm_macro_f1": ds.svm_macro_f1,
                "qsvm_macro_f1": ds.qsvm_macro_f1,
                "accuracy_gap": ds.svm_accuracy - ds.qsvm_accuracy,
                "total_classical_seconds": ds.total_classical_seconds,
                "total_qsvm_seconds": ds.total_qsvm_seconds,
                "gpu_available": ds.gpu_available,
                "gpu_used": ds.gpu_used,
                "confusion_note": ds.confusion_note,
                "best_classical_params": ds.best_classical_params,
                "best_qsvm_params": ds.best_qsvm_params,
                "log_highlights": ds.log_highlights,
            }
            for ds in datasets
        ],
        "portfolio_findings": [
            "Classical SVM matched or exceeded QSVM accuracy on all four datasets.",
            "QSVM search was always the longest phase, making runtime the dominant tradeoff.",
            "Heart Disease was the only run with GPU visibility, but the GPU still appeared mostly idle.",
            "The first Iris attempt failed because the classical SVM grid passed a scalar C instead of a list.",
        ],
        "ranking_by_gap": [ds.display_name for ds in ranked],
    }
    SUMMARY_PATH.write_text(json.dumps(summary, indent=2), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    datasets = [summarize_dataset(key, folder) for key, folder in DATASETS]
    build_summary_json(datasets)
    draw_bar_chart(datasets)
    draw_runtime_chart(datasets)
    draw_log_findings_card(
        [
            "An early Iris run failed because the classical SVM grid used a scalar `C=0.1` instead of a single-item list; later reruns completed cleanly.",
            "QSVM Optuna search logged repeated trial failures (`'NoneType' object has no attribute 'items'`) on Iris, Wine, Heart Disease, and Breast Cancer.",
            "Despite those failed trials, every final experiment run reached `Experiment completed successfully.` and produced full reports, tables, and plots.",
            "Resource summaries show CPU saturation in nearly every major phase, while Heart Disease also reported a mostly idle GPU in several phases.",
        ]
    )
    build_docx(datasets)
    build_markdown(datasets)


if __name__ == "__main__":
    main()
